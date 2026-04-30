import re
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .claude_client import load_prompt, call_claude

OUTPUTS_DIR = Path(__file__).parent.parent / "outputs" / "scripts"


def generate_script(topic: str, fr_angle: str, approved_hook: str) -> dict:
    """
    Returns:
      { full_script, word_count, estimated_mins, docx_path }
    """
    prompt_template = load_prompt("script_prompt")
    prompt = prompt_template.format(
        topic=topic,
        fr_angle=fr_angle,
        approved_hook=approved_hook,
    )
    # Scripts are long — use a higher token limit
    raw_script = call_claude(prompt, max_tokens=8000)

    word_count, estimated_mins = _parse_metadata(raw_script)
    if word_count == 0:
        word_count = len(raw_script.split())
        estimated_mins = round(word_count / 150, 1)  # ~150 words/min narration pace

    slug = _slugify(topic)
    docx_path = _export_docx(topic, raw_script, slug)

    return {
        "full_script": raw_script,
        "word_count": word_count,
        "estimated_mins": estimated_mins,
        "docx_path": str(docx_path),
    }


def _parse_metadata(script: str) -> tuple[int, float]:
    """Extract WORD COUNT and ESTIMATED RUNTIME from script tail."""
    wc_match = re.search(r"WORD COUNT:\s*([0-9,]+)", script)
    rt_match = re.search(r"ESTIMATED RUNTIME:\s*([0-9.]+)", script)
    word_count = int(wc_match.group(1).replace(",", "")) if wc_match else 0
    estimated_mins = float(rt_match.group(1)) if rt_match else 0.0
    return word_count, estimated_mins


def _slugify(text: str) -> str:
    slug = re.sub(r"[^\w\s-]", "", text.lower())
    slug = re.sub(r"[\s_]+", "-", slug).strip("-")
    return slug[:60]


def _export_docx(topic: str, script: str, slug: str) -> Path:
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{slug}-{date.today().isoformat()}.docx"
    path = OUTPUTS_DIR / filename

    doc = Document()

    # Title
    title_para = doc.add_heading(f"FORTUNE & RUIN — {topic.upper()}", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x18, 0x18, 0x18)

    doc.add_paragraph(f"Generated: {date.today().isoformat()}")
    doc.add_paragraph()

    lines = script.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("CHAPTER"):
            # Chapter header
            p = doc.add_heading(line, level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xC0, 0x8A, 0x00)
        elif line.startswith("[") and line.endswith("]"):
            # Tone note or brand bumper — italicised
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif line.startswith("[VISUAL CUE"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

        i += 1

    doc.save(path)
    return path
